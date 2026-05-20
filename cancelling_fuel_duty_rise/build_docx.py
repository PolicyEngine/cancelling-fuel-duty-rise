"""Build the briefing as a Word .docx file with PNG-rendered charts."""

from __future__ import annotations

import argparse
import tempfile
from pathlib import Path

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .charts import (
    annual_cost_chart,
    distributional_chart,
    obr_style_chart,
    rate_path_chart,
)
from .data import compute_all
from .theme import register_template

DEFAULT_OUTPUT = Path(__file__).resolve().parents[1] / "results" / "analysis.docx"

PE_BLUE_RGB = RGBColor(0x2C, 0x64, 0x96)


def _save_fig(fig: go.Figure, path: Path, w: int = 1100, h: int = 550) -> Path:
    fig.write_image(path, width=w, height=h, scale=2)
    return path


def _add_heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Roboto"
    run.font.size = Pt(20 - level * 2)
    run.font.bold = True
    run.font.color.rgb = PE_BLUE_RGB


def _add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Roboto"
    run.font.size = Pt(11)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(11)


def _add_image(doc: Document, path: Path, width_inches: float = 6.2) -> None:
    doc.add_picture(str(path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = PE_BLUE_RGB
    for ri, (_, row) in enumerate(df.iterrows(), start=1):
        for ci, h in enumerate(df.columns):
            v = row[h]
            table.rows[ri].cells[ci].text = (
                f"{v:,.2f}" if isinstance(v, float) else str(v)
            )


def build(output: Path | str = DEFAULT_OUTPUT) -> Path:
    register_template()
    r = compute_all()
    h = r.headline

    with tempfile.TemporaryDirectory(prefix="fueldoc_") as td:
        td = Path(td)
        c1 = _save_fig(annual_cost_chart(r.scrap_5p), td / "c1.png", 900, 460)
        c2 = _save_fig(rate_path_chart(r.rate_path), td / "c2.png", 950, 520)
        c3 = _save_fig(obr_style_chart(r.revenue_2010_2029), td / "c3.png", 1100, 600)
        c_quart = _save_fig(
            distributional_chart(
                r.quartiles,
                group_label="Income quartile (Q1 = lowest, Q4 = highest)",
                title=f"Saving from keeping the 5p cut, by quartile (bottom 5% excluded, {h['year_dist']})",
            ),
            td / "cquart.png",
            900,
            440,
        )
        c_quint = _save_fig(
            distributional_chart(
                r.quintiles,
                group_label="Income quintile (Q1 = lowest, Q5 = highest)",
                title=f"Saving from keeping the 5p cut, by quintile (bottom 5% excluded, {h['year_dist']})",
            ),
            td / "cquint.png",
            900,
            440,
        )
        c_dec = _save_fig(
            distributional_chart(
                r.deciles,
                group_label="Income decile (D1 = next lowest, D10 = highest)",
                title=f"Saving from keeping the 5p cut, by decile (bottom 5% excluded, {h['year_dist']})",
            ),
            td / "cdec.png",
            900,
            440,
        )

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Roboto"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Title
        title = doc.add_paragraph()
        tr = title.add_run("Impact of cancelling the planned fuel duty rise")
        tr.font.size = Pt(22)
        tr.font.bold = True
        tr.font.name = "Roboto"
        tr.font.color.rgb = PE_BLUE_RGB

        # Lede
        lede = doc.add_paragraph()
        lede.paragraph_format.space_after = Pt(14)
        lr = lede.add_run(
            "Rachel Reeves is expected to announce on Thursday that she is shelving the planned 5p "
            "fuel-duty reversal. This briefing reports the cost of extending the cut, the revenue lost "
            "since the first freeze in 2011, the duty rate under an RPI counterfactual, the distributional "
            "impact, and the cross-check against the Guardian and Fleet News figures."
        )
        lr.font.size = Pt(11.5)
        lr.font.italic = True

        # Headline numbers at a glance
        _add_heading(doc, "Headline numbers at a glance", 2)
        kpis = pd.DataFrame(
            [
                [
                    "Cost of cancelling the planned 5p reversal (2027-28)",
                    f"£{h['scrap_2027']:.2f} bn",
                ],
                [
                    "Cost of extending the 5p cut (Guardian framing, 2027-28)",
                    f"£{h['guardian_2027']:.2f} bn",
                ],
                [
                    "Cumulative cost of freezes 2010-11 → 2026-27",
                    f"£{h['fleet_cumulative']:.0f} bn",
                ],
                [
                    "Rate today if uprated by RPI since 2011",
                    f"{h['counterfactual_rate_2026_p']:.1f}p / L (vs actual {h['actual_rate_2026_p']:.1f}p)",
                ],
                [
                    f"Annual revenue gap by {h['last_year']}-{(h['last_year'] + 1) % 100:02d} vs counterfactual",
                    f"£{h['revenue_last_year_counterfactual_bn'] - h['revenue_last_year_actual_bn']:.0f} bn",
                ],
            ],
            columns=["Metric", "Value"],
        )
        _add_table(doc, kpis)

        # Story so far
        _add_heading(doc, "The story so far", 2)
        _add_para(
            doc,
            "Fuel duty is the per-litre tax on petrol and diesel at the pump. The rate has been frozen every "
            "year since Budget 2011. In March 2022 Rishi Sunak cut it by a further 5p (from 57.95p to 52.95p) "
            "as a temporary measure. The Autumn Budget 2025 set out a phased reversal of that cut — 1p in "
            "September 2026, 2p in December 2026, 2p in March 2027 — followed by annual RPI uprating from "
            "April 2027 onwards. Reeves is expected to cancel the September 1p step-up and may cancel all of "
            "the 5p reversal; the press reports do not address the April-2027 RPI uprating.",
        )

        _add_heading(doc, "How much does scrapping the 5p increase cost?", 2)
        _add_para(
            doc,
            "The Autumn Budget 2025 schedule increases the duty by 1p in September 2026, 2p in December "
            "2026, 2p in March 2027, then uprates by RPI each April. Holding the rate at 52.95p instead, "
            "the annual revenue forgone is:",
        )
        _add_image(doc, c1)
        _add_table(
            doc,
            r.scrap_5p[r.scrap_5p["Year"] >= 2026].assign(
                Year=lambda d: d["Year"].astype(str)
            ),
        )

        _add_heading(doc, "What would the rate be if it had risen with RPI since 2011?", 2)
        _add_para(
            doc,
            "Compounding RPI annually onto the 57.95p rate frozen in Budget 2011 gives the counterfactual "
            "path below.",
        )
        _add_image(doc, c2, 6.4)

        _add_heading(doc, "How much money have the freezes lost so far?", 2)
        _add_para(doc, "The chart combines three series:")
        _add_bullet(
            doc,
            "Yellow line — HMRC out-turn (2010-11 to 2024-25): fuel-duty receipts as published in HMRC's "
            "UK Tax & NICs receipts statistics on gov.uk.",
        )
        _add_bullet(
            doc,
            "Blue line — PolicyEngine UK projection (2025-26 onwards): revenue from running the "
            "PolicyEngine UK microsimulation on the enhanced FRS 2023-29 dataset at the current-law rate "
            "schedule.",
        )
        _add_bullet(
            doc,
            "Dashed teal line — RPI counterfactual: each year's revenue scaled by counterfactual_rate / "
            "actual_rate. Fuel volumes are held fixed (no behavioural response).",
        )
        _add_image(doc, c3, 6.6)

        _add_heading(doc, "Does this match the Guardian and Fleet News?", 2)
        _add_para(
            doc,
            f"The two press reports on 18 May 2026 quoted £2.4 bn / year (Guardian) and ~£120 bn "
            f"cumulative since 2010/11 (Fleet News). Both use the 'extend the 5p cut' framing: 52.95p kept "
            f"versus a return to 57.95p, with no further RPI uprating. In 2027-28, PolicyEngine UK puts "
            f"that figure at £{h['guardian_2027']:.2f} bn (and £{h['guardian_2026']:.2f} bn in 2026-27), "
            f"while the cumulative cost of freezes from 2010/11 to 2026/27 comes to "
            f"£{h['fleet_cumulative']:.1f} bn — both consistent with the press numbers. The earlier "
            f"section reports a higher 2027-28 figure (£{h['scrap_2027']:.2f} bn) because it compares "
            f"against 59.25p — i.e. 57.95p plus the April-2027 RPI uprating. The £"
            f"{h['scrap_2027'] - h['guardian_2027']:.2f} bn difference is the cost of cancelling that RPI "
            f"uprating on top of the 5p reversal.",
        )
        _add_para(
            doc,
            "Guardian article: https://www.theguardian.com/politics/2026/may/18/rachel-reeves-fuel-duty-cost-of-living",
        )

        _add_heading(doc, "Who gains from keeping the cut?", 2)
        _add_para(
            doc,
            "Average saving per household if the 5p cut is kept, as a share of household net income. "
            "The bottom 5% by equivalised income is excluded from all three cuts (Resolution Foundation "
            "approach). The remaining 95% is then split into quartiles, quintiles and deciles.",
        )

        _add_heading(doc, "By income quartile", 3)
        _add_image(doc, c_quart)
        _add_table(doc, r.quartiles)

        _add_heading(doc, "By income quintile", 3)
        _add_image(doc, c_quint)
        _add_table(doc, r.quintiles)

        _add_heading(doc, "By income decile", 3)
        _add_image(doc, c_dec)
        _add_table(doc, r.deciles)

        _add_heading(doc, "Sources", 2)
        _add_bullet(doc, f"Household-level figures: {r.citation}.")
        _add_bullet(
            doc,
            "Historical fuel-duty receipts (2010-11 → 2024-25): HMRC UK Tax & NICs receipts publication "
            "(gov.uk).",
        )
        _add_bullet(doc, "RPI series: OBR Economic and Fiscal Outlook, March 2026.")
        _add_bullet(
            doc,
            "No behavioural responses modelled: fuel volumes held fixed across scenarios.",
        )

        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", "-o", default=DEFAULT_OUTPUT, type=Path)
    args = parser.parse_args()
    path = build(args.output)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
